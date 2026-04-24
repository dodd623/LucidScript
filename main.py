from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse, RedirectResponse
from starlette.middleware.sessions import SessionMiddleware
from pydantic import BaseModel
import asyncio
import tempfile, os, pathlib, uuid, re, subprocess, shlex, textwrap, html
from typing import List, Tuple, Optional
import whisper
import easyocr
import json
from passlib.context import CryptContext
from deep_translator import GoogleTranslator
from docx import Document
from datetime import datetime
from sqlalchemy import create_engine, Column, Integer, String, Boolean, DateTime, Text
from sqlalchemy.orm import declarative_base, sessionmaker

app = FastAPI()

SECRET_KEY = os.getenv("SECRET_KEY", "lucidscript-dev-secret-change-me")
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY)

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

APP_VERSION = os.getenv("APP_VERSION", "0.3.1").strip()

model_name = os.getenv("WHISPER_MODEL", "medium").strip().lower()
allowed_models = {"tiny", "base", "small", "medium", "large"}
if model_name not in allowed_models:
    model_name = "medium"

model = None


def get_model():
    global model
    if model is None:
        model = whisper.load_model(model_name)
    return model


ocr_reader = None
ocr_reader_ch = None
ocr_reader_ja = None

BASE_DIR = pathlib.Path(__file__).parent.resolve()
OUTPUT_DIR = (BASE_DIR / "output").resolve()
OUTPUT_DIR.mkdir(exist_ok=True)

DB_PATH = BASE_DIR / "lucidscript.db"
DATABASE_URL = f"sqlite:///{DB_PATH.as_posix()}"

engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine, autoflush=False, autocommit=False)
Base = declarative_base()

TEMPLATE_DIR = (BASE_DIR / "templates").resolve()
WITNESS_TEMPLATE_PATH = TEMPLATE_DIR / "witness_statement_template.docx"

UI_TEMPLATE_PATH = BASE_DIR / "frontend" / "ui.html"
try:
    UI_TEMPLATE = UI_TEMPLATE_PATH.read_text(encoding="utf-8")
except FileNotFoundError:
    UI_TEMPLATE = None


class User(Base):
    __tablename__ = "users"

    id = Column(Integer, primary_key=True, index=True)
    username = Column(String, unique=True, nullable=False, index=True)
    email = Column(String, unique=True, nullable=False, index=True)
    password_hash = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow, nullable=False)


class DocumentRecord(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    created_at = Column(DateTime, default=datetime.utcnow, nullable=False)

    user_id = Column(Integer, nullable=True, index=True)

    mode = Column(String, nullable=False)
    original_filename = Column(String, nullable=True)
    output_filename = Column(String, nullable=False)

    status = Column(String, nullable=False, default="completed")
    language = Column(String, nullable=True)
    translated = Column(Boolean, default=False)

    error_message = Column(Text, nullable=True)
    notes = Column(Text, nullable=True)


Base.metadata.create_all(bind=engine)


def save_document_record(
    mode: str,
    output_filename: str,
    original_filename: str | None = None,
    status: str = "completed",
    language: str | None = None,
    translated: bool = False,
    error_message: str | None = None,
    notes: str | None = None,
    user_id: int | None = None,
):
    db = SessionLocal()
    try:
        record = DocumentRecord(
            user_id=user_id,
            mode=mode,
            original_filename=original_filename,
            output_filename=output_filename,
            status=status,
            language=language,
            translated=translated,
            error_message=error_message,
            notes=notes,
        )
        db.add(record)
        db.commit()
        db.refresh(record)
        return record
    finally:
        db.close()


@app.get("/documents")
def list_documents():
    db = SessionLocal()
    try:
        records = (
            db.query(DocumentRecord).order_by(DocumentRecord.created_at.desc()).all()
        )
        return [
            {
                "id": r.id,
                "created_at": r.created_at.isoformat() if r.created_at else None,
                "mode": r.mode,
                "original_filename": r.original_filename,
                "output_filename": r.output_filename,
                "status": r.status,
                "language": r.language,
                "translated": r.translated,
                "error_message": r.error_message,
                "notes": r.notes,
            }
            for r in records
        ]
    finally:
        db.close()


def get_ocr_readers():
    global ocr_reader, ocr_reader_ch, ocr_reader_ja

    if ocr_reader is None:
        ocr_reader = easyocr.Reader(["en", "es", "fr", "de", "pt", "it", "nl"])

    if ocr_reader_ch is None:
        ocr_reader_ch = easyocr.Reader(["ch_sim", "en"])

    if ocr_reader_ja is None:
        ocr_reader_ja = easyocr.Reader(["ja", "en"])

    return ocr_reader, ocr_reader_ch, ocr_reader_ja


def landing_page_html() -> str:
    return f"""
    <html data-theme="dark">
      <head>
        <title>LucidScript</title>
        <meta name="viewport" content="width=device-width, initial-scale=1" />
        <style>
          :root {{
            --bg: #0f1115;
            --text: #eaeef3;
            --muted: rgba(234, 238, 243, 0.9);
            --muted-soft: rgba(234, 238, 243, 0.75);
            --card: #171a21;
            --subcard: #11141b;
            --border: #232736;
            --code-bg: #0b0d12;
            --button: #4c83ff;
            --button-hover: #3a6ef6;
            --shadow: none;
          }}

          html[data-theme="light"] {{
            --bg: #f5f7fb;
            --text: #1a2233;
            --muted: rgba(26, 34, 51, 0.88);
            --muted-soft: rgba(26, 34, 51, 0.72);
            --card: #ffffff;
            --subcard: #f7f9fd;
            --border: #d9e1ef;
            --code-bg: #eef3fb;
            --button: #3f6fe5;
            --button-hover: #315bc0;
            --shadow: 0 10px 28px rgba(23, 32, 56, 0.08);
          }}

          * {{ box-sizing: border-box; }}

          body {{
            margin: 0;
            padding: 0;
            font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial;
            background: var(--bg);
            color: var(--text);
            display: flex;
            min-height: 100vh;
            transition: background 0.2s ease, color 0.2s ease;
          }}

          .wrap {{
            margin: auto;
            width: min(920px, 92%);
            text-align: center;
            padding: 32px 0;
          }}

          .topbar {{
            display: flex;
            justify-content: flex-end;
            margin-bottom: 14px;
          }}

          .theme-toggle {{
            border: 1px solid var(--border);
            background: var(--card);
            color: var(--text);
            border-radius: 999px;
            padding: 8px 12px;
            font-size: 13px;
            cursor: pointer;
            box-shadow: var(--shadow);
          }}

          .theme-toggle:hover {{
            opacity: 0.92;
          }}

          h1 {{
            font-weight: 700;
            font-size: clamp(2.2rem, 4vw, 3.4rem);
            margin-bottom: 0.35rem;
          }}

          .sub {{
            color: var(--muted);
            font-size: 1.08rem;
            margin-top: 0.2rem;
            margin-bottom: 1.35rem;
          }}

          .version {{
            display: inline-block;
            margin-top: 2px;
            margin-bottom: 14px;
            padding: 6px 10px;
            border-radius: 999px;
            background: var(--card);
            border: 1px solid var(--border);
            font-size: 12px;
            color: var(--muted-soft);
            box-shadow: var(--shadow);
          }}

          .card {{
            background: var(--card);
            border: 1px solid var(--border);
            border-radius: 16px;
            padding: 28px;
            text-align: left;
            margin-top: 18px;
            box-shadow: var(--shadow);
          }}

          .hero-actions {{
            text-align: center;
            margin: 20px 0 10px 0;
          }}

          a.button {{
            display: inline-block;
            padding: 12px 20px;
            border-radius: 10px;
            background: var(--button);
            color: white;
            font-weight: 600;
            text-decoration: none;
          }}

          a.button:hover {{
            background: var(--button-hover);
          }}

          h2 {{
            margin-top: 0;
            margin-bottom: 10px;
            font-size: 1.15rem;
          }}

          p {{
            color: var(--muted);
            line-height: 1.55;
            margin: 0 0 12px 0;
          }}

          ul, ol {{
            margin: 10px 0 0 18px;
            padding: 0;
            color: var(--muted);
            line-height: 1.6;
          }}

          li {{ margin-bottom: 8px; }}

          .grid {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 16px;
            margin-top: 16px;
          }}

          .mini {{
            background: var(--subcard);
            border: 1px solid var(--border);
            border-radius: 14px;
            padding: 18px;
          }}
          
          .mini h3 {{
            margin-top: 12px;
            margin-bottom: 6px;
            font-size: 14px;
            opacity: 0.85;
          }}

          .hint {{
            margin-top: 16px;
            font-size: 12px;
            color: var(--muted-soft);
            text-align: center;
          }}

          code {{
            background: var(--code-bg);
            padding: 2px 6px;
            border-radius: 6px;
          }}

          @media (max-width: 700px) {{
            .grid {{ grid-template-columns: 1fr; }}
            .card {{ padding: 22px; }}
          }}
        </style>
      </head>
      <body>
        <div class="wrap">
          <div class="topbar">
            <button class="theme-toggle" id="theme-toggle" type="button">Toggle theme</button>
          </div>

          <h1>LucidScript</h1>
          <div class="sub">
            A document-generation assistant for transcription, Optical Character Recognition (OCR), and incident-style report workflows.
          </div>
          <div class="version">Version {html.escape(APP_VERSION)} • Whisper model: {html.escape(model_name)}</div>

          <div class="hero-actions">
            <a href="/ui_async" class="button">Open LucidScript UI</a>
          </div>

          <div class="card">
            <h2>What LucidScript does</h2>
            <p>
              LucidScript is designed to take raw source material like audio, typed notes, image-based text,
              and other video content, then turn that information into clean, downloadable Word documents.
            </p>
            <p>
              The goal is to reduce manual formatting work and make it easier to move from rough evidence or notes
              to polished documentation. In plain English: less copy-paste goblin labor, more usable reports.
            </p>

            <div class="grid">
              <div class="mini">
                <h2>Current modes</h2>
                <ul>
                  <li><strong>Audio Transcription:</strong> upload supported audio/video files and export transcript-style documents.</li>
                  <li><strong>Text Input:</strong> paste written content and convert it into a formatted report document.</li>
                  <li><strong>Image Upload:</strong> extract text from images and optionally translate it to English before export.</li>
                </ul>
              </div>

              <div class="mini">
                <h2>Why it matters</h2>
                <ul>
                  <li>Speeds up documentation workflows.</li>
                  <li>Supports converting raw inputs into cleaner reports.</li>
                  <li>Helps bridge transcription, OCR, and reporting in one place.</li>
                  <li>Creates downloadable <code>.docx</code> outputs for practical use.</li>
                </ul>
              </div>
            </div>

            <div class="grid">
              <div class="mini">
                <h2>Audio processing pipeline</h2>
                <br/>
                <ol>          
                  <li><strong>Upload:</strong> the user uploads an audio or video file through the interface.</li>
                  <li><strong>Validation:</strong> the system checks the file extension and accepts supported media types.</li>
                  <li><strong>Temporary staging:</strong> the upload is saved as a temporary local file for processing.</li>
                  <li><strong>Transcription:</strong> the file is processed with <strong>OpenAI Whisper</strong> using the configured model.</li>
                  <li><strong>Optional translation:</strong> if selected, Whisper runs in translate mode and outputs English text.</li>
                  <li><strong>Optional speaker flow:</strong> if deposition mode and diarization are selected, the audio is converted to 16k WAV and speaker labels are assigned.</li>
                  <li><strong>Formatting:</strong> the transcript is organized into either standard paragraphs or deposition-style speaker sections.</li>
                  <li><strong>DOCX generation:</strong> a Word document is created and saved to the output folder.</li>
                  <li><strong>Download:</strong> the final file is returned to the user through a download link.</li>
                  <br/>
                  <br/>
                  <br/>
                  <br/>
                  <h2>AI model explanation</h2>
                  <p>
                    LucidScript uses a combination of locally running machine learning models and lightweight translation tools
                    to process different types of input data.
                  </p>

                  <ul>
                    <li>
                     <strong>Whisper (OpenAI):</strong>
                      Used for audio transcription. Whisper converts spoken audio into text and can optionally translate
                      non-English speech into English. It runs locally using the configured model size
                      (<code>{html.escape(model_name)}</code>), balancing speed and accuracy.
                    </li>

                    <li>
                      <strong>EasyOCR:</strong>
                        Used for extracting text from images. Multiple OCR readers are initialized to support
                        different languages (including English, Chinese, and Japanese), with fallback logic
                        if the primary reader fails.
                    </li>

                      <li>
                        <strong>GoogleTranslator (deep-translator):</strong>
                        Used for optional translation of extracted or transcribed text into English.
                        This step is only applied when the user selects translation.
                      </li>
                  </ul>
                  <br/>
                  <p>
                    These models were chosen to allow LucidScript to run without requiring heavy external APIs,
                    making the system more portable and easier to deploy while still supporting multilingual input.
                  </p>
                  
                </ol>
              </div>

              <div class="mini">               
                <h2>Project overview</h2>
                <br/>
                <h3>Direction</h3>
                <p>
                  The broader vision is to make LucidScript useful for documentation-heavy environments where users
                  may need to process interviews, notes, images, or statement data quickly and consistently.
                </p>
                <br/>
                <h2>API usage</h2>
                  <p>
                    LucidScript currently relies mostly on locally running tools instead of heavy paid external APIs.
                    Whisper transcription runs locally on the server using the configured model, and EasyOCR also runs
                    locally for image text extraction. This reduces external dependencies and keeps the system portable and easier to deploy.
                  </p>
                <p>
                  The main external-style service in the current workflow is translation through
                  <code>deep-translator</code>. That is only used when the user selects translation. In a larger shared
                  production version, API usage would need rate limits, key management, usage monitoring, and a plan for
                  handling costs if multiple users were submitting jobs at the same time.
                </p>
                <br/>
                <h2>Technical risks</h2>
                  <ul>
                    <li><strong>Audio processing load:</strong> Whisper transcription is CPU-intensive and can slow health checks or response times during longer uploads.</li>
                    <li><strong>OCR reliability:</strong> image extraction quality depends heavily on resolution, lighting, layout, and language support.</li>
                    <li><strong>Translation dependency:</strong> optional translation adds another failure point if the translation service is unavailable or inconsistent.</li>
                    <li><strong>Pipeline complexity:</strong> audio, OCR, videos, and formatted exports each fail differently, so debugging must identify the exact step where processing broke.</li>
                    <li><strong>Storage growth:</strong> generated output files and saved document history will continue growing unless retention or cleanup rules are added.</li>
                    <li><strong>Deployment limits:</strong> lightweight hosting is fine for the MVP, but larger workloads may require stronger compute resources or background job handling.</li>
                  </ul>
                  <br/>
                  <h2>Helpful links</h2>
                  <ul>
                    <li>Main interface: <code>/ui_async</code></li>
                    <li>Health check: <code>/health</code></li>
                    <li>API docs: <code>/docs</code></li>
                    <li>Downloads: <code>/download/&lt;filename&gt;.docx</code></li>
                  </ul>
                  <h2>Version info</h2>
                  <ul>
                    <li><strong>Build:</strong> 0.3.1</li>
                    <li><strong>Whisper model:</strong> {html.escape(model_name)}</li>
                    <li>Version is displayed across UI and endpoints for debugging.</li>
                  </ul>
              </div>
                

        <script>
          (function () {{
            const root = document.documentElement;
            const toggle = document.getElementById("theme-toggle");
            const saved = localStorage.getItem("lucidscript-theme");
            const preferred = saved || "dark";

            function applyTheme(theme) {{
              root.setAttribute("data-theme", theme);
              localStorage.setItem("lucidscript-theme", theme);
              toggle.textContent = theme === "dark" ? "Switch to light mode" : "Switch to dark mode";
            }}

            applyTheme(preferred);

            toggle.addEventListener("click", function () {{
              const current = root.getAttribute("data-theme") || "dark";
              applyTheme(current === "dark" ? "light" : "dark");
            }});
          }})();
        </script>
      </body>
    </html>
    """


@app.get("/", response_class=HTMLResponse)
async def root():
    return landing_page_html()


@app.get("/health")
async def health_check():
    return {
        "status": "ok",
        "whisper_model": model_name,
        "version": APP_VERSION,
    }


@app.get("/test_youtube")
def test_youtube(url: str):
    path, meta = download_youtube_audio(url)

    return {
        "audio_path": path,
        "title": meta.get("title"),
        "duration": meta.get("duration"),
    }


@app.get("/ui", response_class=HTMLResponse)
def upload_ui():
    return """
    <html>
      <head>
        <title>LucidScript</title>
        <meta name="viewport" content="width=device-width, initial-scale=1" />
        <style>
          :root { color-scheme: dark; }
          body { margin:0; padding:0; font-family:system-ui, -apple-system, Segoe UI, Roboto, Arial;
                 background:#0f1115; color:#eaeef3; display:flex; min-height:100vh; }
          .wrap { margin:auto; width:min(640px, 92%); text-align:center; }
          h1 { font-weight:700; letter-spacing:.3px; margin-bottom:.2rem; }
          p { opacity:.8; margin-top:.2rem; margin-bottom:1.2rem; }
          .card { background:#171a21; border:1px solid #232736; border-radius:14px; padding:24px; }
          input[type=file] { width:100%; background:#0f1115; color:#eaeef3; border:1px dashed #2a3042;
                             padding:14px; border-radius:10px; }
          button { margin-top:14px; width:100%; padding:12px 16px; border:0; border-radius:10px;
                   background:#4c83ff; color:white; font-weight:600; cursor:pointer; }
          button:hover { background:#3a6ef6; }
          small { display:block; margin-top:10px; opacity:.65; }
          a { color:#9ec1ff; text-decoration:none; }
          .hint { margin-top:10px; font-size:12px; opacity:.8; }
          code { background:#0b0d12; padding:2px 6px; border-radius:6px; }
        </style>
      </head>
      <body>
        <div class="wrap">
          <h1>LucidScript</h1>
         <p>Upload audio or paste a YouTube URL → optional language/translate → choose output style → download .docx.</p>
          <div class="card">
            <form action="/export_docx_from_audio" enctype="multipart/form-data" method="post">
              <label for="file">Select an audio file:</label><br/><br/>
              <input
                id="file"
                type="file"
                name="file"
                accept=".wav,.mp3,.m4a,.aac,.flac,.ogg,.webm,.mp4,audio/*,video/*"
                required
              />
              <div class="hint">
                Supported: <code>WAV</code>, <code>MP3</code>, <code>M4A</code>, <code>AAC</code>, <code>FLAC</code>, <code>OGG</code>, <code>WEBM</code>, <code>MP4</code>
              </div>
              <button type="submit">Transcribe & Export</button>
            </form>
            <small>Prefer the API? Try <a href="/docs">/docs</a> or the async UI at <a href="/ui_async">/ui_async</a>.</small>
          </div>
        </div>
      </body>
    </html>
    """


class FormatRequest(BaseModel):
    raw_text: str


def to_paragraphs(text: str):
    text = re.sub(r"\s+", " ", text).strip()
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", text)
    return [p.strip() for p in parts if p.strip()]


def build_transcript_doc(
    title: str,
    text: str,
    language: str | None = None,
    translated: bool = False,
):
    doc = Document()
    doc.add_heading(title, 0)

    meta = datetime.now().strftime("%Y-%m-%d %H:%M")
    if language:
        meta += f"  |  Language: {language}"
    if translated:
        meta += "  |  Translated to English"

    doc.add_paragraph(meta)

    for p in to_paragraphs(text):
        doc.add_paragraph(p)

    out = OUTPUT_DIR / f"lucidscript_{uuid.uuid4().hex[:8]}.docx"
    doc.save(out.as_posix())
    return out


def build_security_report_doc(text: str):
    doc = Document()
    doc.add_heading("Security Report", 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))

    for block in text.splitlines():
        cleaned = block.strip()
        if cleaned:
            doc.add_paragraph(cleaned)
        else:
            doc.add_paragraph("")

    out = OUTPUT_DIR / f"text_transcript_{uuid.uuid4().hex[:8]}.docx"
    doc.save(out.as_posix())
    return out


def build_multi_image_ocr_doc(
    file_results: List[dict],
    title: str = "LucidScript OCR Extraction",
):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M"))
    doc.add_paragraph(f"Image count: {len(file_results)}")

    for index, item in enumerate(file_results, start=1):
        filename = item["filename"]
        extracted_text = (item.get("final_text") or "").strip()

        doc.add_heading(f"Image {index}: {filename}", level=1)

        if extracted_text:
            for block in extracted_text.splitlines():
                cleaned = block.strip()
                if cleaned:
                    doc.add_paragraph(cleaned)
                else:
                    doc.add_paragraph("")
        else:
            doc.add_paragraph("[No readable text detected]")

    out = OUTPUT_DIR / f"ocr_extraction_{uuid.uuid4().hex[:8]}.docx"
    doc.save(out.as_posix())
    return out


def replace_placeholders_in_paragraph(paragraph, replacements: dict):
    full_text = "".join(run.text for run in paragraph.runs)
    if not full_text:
        return

    updated_text = full_text
    for key, value in replacements.items():
        updated_text = updated_text.replace(key, value)

    if updated_text != full_text:
        if paragraph.runs:
            paragraph.runs[0].text = updated_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = updated_text


def replace_placeholders_in_doc(doc: Document, replacements: dict):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)


# def build_witness_statement_doc(
#     witness_name: str,
#     occupation: str,
#     statement_body: str,
#     age_text: str = "Over 18",
#     date_text: str | None = None,
# ):
#     if not WITNESS_TEMPLATE_PATH.exists():
#         raise HTTPException(
#             status_code=500,
#             detail="Witness statement template file was not found in the templates folder.",
#         )

#     doc = Document(WITNESS_TEMPLATE_PATH.as_posix())

#     replacements = {
#         "{{WITNESS_NAME}}": witness_name.strip(),
#         "{{OCCUPATION}}": occupation.strip(),
#         "{{STATEMENT_BODY}}": statement_body.strip(),
#         "{{DATE}}": (date_text or datetime.now().strftime("%m/%d/%Y")).strip(),
#         "{{AGE}}": age_text.strip(),
#     }

#     replace_placeholders_in_doc(doc, replacements)

#     out = OUTPUT_DIR / f"witness_statement_{uuid.uuid4().hex[:8]}.docx"
#     doc.save(out.as_posix())
#     return out


def extract_text_from_image(image_path: str) -> str:
    local_ocr_reader, local_ocr_reader_ch, local_ocr_reader_ja = get_ocr_readers()

    try:
        results = local_ocr_reader.readtext(image_path, detail=0, paragraph=True)
        text = "\n".join(
            [line.strip() for line in results if line and line.strip()]
        ).strip()
        if text:
            return text
    except Exception:
        pass

    try:
        results_ch = local_ocr_reader_ch.readtext(image_path, detail=0, paragraph=True)
        text_ch = "\n".join(
            [line.strip() for line in results_ch if line and line.strip()]
        ).strip()
        if text_ch:
            return text_ch
    except Exception:
        pass

    try:
        results_ja = local_ocr_reader_ja.readtext(image_path, detail=0, paragraph=True)
        text_ja = "\n".join(
            [line.strip() for line in results_ja if line and line.strip()]
        ).strip()
        if text_ja:
            return text_ja
    except Exception:
        pass

    return ""


def translate_text_to_english(text: str) -> str:
    if not text.strip():
        return text

    try:
        return GoogleTranslator(source="auto", target="en").translate(text)
    except Exception:
        return text


def translate_mixed_text_to_english(text: str) -> str:
    """
    Translates OCR text line-by-line so bilingual/trilingual images keep their structure.
    English lines usually remain unchanged, while non-English lines are translated.
    """
    if not text.strip():
        return text

    translated_lines = []

    for line in text.splitlines():
        cleaned = line.strip()

        if not cleaned:
            translated_lines.append("")
            continue

        try:
            translated = GoogleTranslator(source="auto", target="en").translate(cleaned)
            translated_lines.append(translated or cleaned)
        except Exception:
            translated_lines.append(cleaned)

    return "\n".join(translated_lines).strip()


@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    suffix = os.path.splitext(file.filename or "")[-1]
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name
        result = await asyncio.to_thread(get_model().transcribe, tmp_path)
        text = (result.get("text") or "").strip()
        return {"transcript": text}
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Something went wrong while processing the audio. Please try again with a different file.",
        )
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass


@app.post("/format_docx")
def format_docx(req: FormatRequest):
    if not req.raw_text.strip():
        raise HTTPException(
            status_code=400,
            detail="No text was provided to format.",
        )
    out = build_transcript_doc("LucidScript Transcript", req.raw_text)
    return {"docx_path": str(out)}


# @app.post("/export_security_report")
# async def export_security_report(report_text: str = Form(...)):
#     if not report_text.strip():
#         raise HTTPException(
#             status_code=400,
#             detail="No report text was provided.",
#         )

#     out = await asyncio.to_thread(build_security_report_doc, report_text)

#     save_document_record(
#         mode="text",
#         original_filename=None,
#         output_filename=out.name,
#         status="completed",
#         notes="Generated from pasted text input.",
#     )

#     return JSONResponse(
#         {
#             "message": "Text transcript formatted successfully.",
#             "docx_path": str(out),
#             "docx_filename": out.name,
#         }
#     )


@app.post("/export_multi_image_ocr")
async def export_multi_image_ocr(
    image_files: List[UploadFile] = File(...),
    translate_to_english: str | None = Form(None),
):
    allowed = {".png", ".jpg", ".jpeg", ".webp", ".bmp"}
    translated_flag = (translate_to_english or "").lower() == "true"

    if not image_files:
        raise HTTPException(
            status_code=400,
            detail="At least one image file is required.",
        )

    tmp_paths: List[str] = []
    file_results: List[dict] = []

    try:
        for image_file in image_files:
            suffix = os.path.splitext(image_file.filename or "")[-1].lower()

            if suffix not in allowed:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported image format for {image_file.filename or 'uploaded file'}. Use PNG, JPG, JPEG, WEBP, or BMP.",
                )

            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(await image_file.read())
                tmp_path = tmp.name
                tmp_paths.append(tmp_path)

            print(f"[OCR] Processing: {image_file.filename}")
            extracted_text = await asyncio.to_thread(extract_text_from_image, tmp_path)
            print(f"[OCR] Extracted ({image_file.filename}): {extracted_text[:100]}")
            final_text = (
              await asyncio.to_thread(translate_mixed_text_to_english, extracted_text)
              if translated_flag
              else extracted_text
            )

            file_results.append(
                {
                    "filename": image_file.filename or pathlib.Path(tmp_path).name,
                    "extracted_text": extracted_text,
                    "final_text": final_text,
                    "has_text": bool(extracted_text.strip()),
                }
            )

        if not any(item["has_text"] for item in file_results):
            print("[OCR] No readable text found in any images, continuing anyway.")

        out = await asyncio.to_thread(
            build_multi_image_ocr_doc,
            file_results,
            "LucidScript OCR Extraction",
        )

        save_document_record(
            mode="image",
            original_filename=", ".join(item["filename"] for item in file_results),
            output_filename=out.name,
            status="completed",
            translated=translated_flag,
            notes="Generated from generic multi-image OCR workflow.",
        )

        return JSONResponse(
            {
                "message": "OCR extraction generated successfully.",
                "docx_path": str(out),
                "docx_filename": out.name,
                "files": file_results,
                "translated": translated_flag,
                "image_count": len(file_results),
                "version": APP_VERSION,
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback

        print("[OCR] ERROR:")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"OCR route failed: {str(e)}",
        )
    finally:
        for tmp_path in tmp_paths:
            try:
                if tmp_path and os.path.exists(tmp_path):
                    os.remove(tmp_path)
            except Exception:
                pass


# @app.post("/export_witness_statement")
# async def export_witness_statement(
#     witness_name: str = Form(...),
#     occupation: str = Form(...),
#     statement_body: str = Form(...),
#     age_text: str = Form("Over 18"),
# ):
#     if not witness_name.strip():
#         raise HTTPException(status_code=400, detail="Witness name is required.")
#     if not occupation.strip():
#         raise HTTPException(status_code=400, detail="Occupation is required.")
#     if not statement_body.strip():
#         raise HTTPException(status_code=400, detail="Statement body is required.")

#     out = await asyncio.to_thread(
#         build_witness_statement_doc,
#         witness_name=witness_name,
#         occupation=occupation,
#         statement_body=statement_body,
#         age_text=age_text,
#     )

#     save_document_record(
#         mode="witness",
#         original_filename=None,
#         output_filename=out.name,
#         status="completed",
#         notes=f"Witness statement for {witness_name.strip()}.",
#     )

#     return JSONResponse(
#         {
#             "message": "Witness statement generated successfully.",
#             "docx_path": str(out),
#             "docx_filename": out.name,
#         }
#     )


@app.post("/export_docx_from_audio")
async def export_docx_from_audio(file: UploadFile = File(...)):
    suffix = os.path.splitext(file.filename or "")[-1]
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        result = await asyncio.to_thread(get_model().transcribe, tmp_path)
        text = (result.get("text") or "").strip()
        if not text:
            raise HTTPException(
                status_code=400,
                detail="No speech was detected in this file. Try a different recording.",
            )

        language = result.get("language", "unknown")
        duration = (
            round(float(result.get("duration", 0)), 2) if "duration" in result else None
        )

        out = await asyncio.to_thread(
            build_transcript_doc,
            "LucidScript Transcript",
            text,
            language,
        )

        save_document_record(
            mode="audio",
            original_filename=file.filename,
            output_filename=out.name,
            status="completed",
            language=language,
            translated=False,
            notes="Generated from basic audio transcription route.",
        )

        return {
            "message": "Transcription finished. Word doc ready to download.",
            "docx_path": str(out),
            "language": language,
            "duration_sec": duration,
        }
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Something went wrong while processing the audio. Please try again with a different file.",
        )
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass


@app.get("/ui_async", response_class=HTMLResponse)
def upload_ui_async():
    page = """
    <html data-theme="dark">
      <head>
        <title>LucidScript — Async</title>
        <meta name="viewport" content="width=device-width, initial-scale=1" />
        <style>
          :root {
            --bg: #0f1115;
            --text: #eaeef3;
            --muted: rgba(234, 238, 243, 0.85);
            --muted-soft: rgba(234, 238, 243, 0.78);
            --card: #171a21;
            --subcard: #11141b;
            --border: #232736;
            --input-bg: #0f1115;
            --code-bg: #0b0d12;
            --button: #4c83ff;
            --button-hover: #3a6ef6;
            --success: #71eea0;
            --error: #ff8a8a;
            --shadow: none;
          }

          html[data-theme="light"] {
            --bg: #f5f7fb;
            --text: #1a2233;
            --muted: rgba(26, 34, 51, 0.86);
            --muted-soft: rgba(26, 34, 51, 0.72);
            --card: #ffffff;
            --subcard: #f7f9fd;
            --border: #d9e1ef;
            --input-bg: #ffffff;
            --code-bg: #eef3fb;
            --button: #3f6fe5;
            --button-hover: #315bc0;
            --success: #1b9b52;
            --error: #d64045;
            --shadow: 0 10px 28px rgba(23, 32, 56, 0.08);
          }

          * { box-sizing: border-box; }

          body {
            margin: 0;
            padding: 0;
            font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial;
            background: var(--bg);
            color: var(--text);
            display: flex;
            min-height: 100vh;
            transition: background 0.2s ease, color 0.2s ease;
          }

          .wrap {
            margin: auto;
            width: min(900px, 94%);
            padding: 32px 0;
          }

          .topbar {
            display: flex;
            justify-content: flex-end;
            margin-bottom: 14px;
          }

          .theme-toggle {
            border: 1px solid var(--border);
            background: var(--card);
            color: var(--text);
            border-radius: 999px;
            padding: 8px 12px;
            font-size: 13px;
            cursor: pointer;
            box-shadow: var(--shadow);
          }

          .theme-toggle:hover {
            opacity: 0.92;
          }

          h1 {
            font-weight: 700;
            letter-spacing: 0.3px;
            margin-bottom: 0.25rem;
          }

          h2 {
            margin-top: 0;
            margin-bottom: 0.5rem;
          }

          h3 {
            margin-top: 0;
            margin-bottom: 0.5rem;
            font-size: 1rem;
          }

          p {
            color: var(--muted);
            margin-top: 0.2rem;
            margin-bottom: 1rem;
          }

          .version {
            display: inline-block;
            margin-top: 0;
            margin-bottom: 14px;
            padding: 6px 10px;
            border-radius: 999px;
            background: var(--card);
            border: 1px solid var(--border);
            font-size: 12px;
            color: var(--muted-soft);
            box-shadow: var(--shadow);
          }

          .card {
            background: var(--card);
            border: 1px solid var(--border);
            border-radius: 14px;
            padding: 24px;
            margin-bottom: 18px;
            box-shadow: var(--shadow);
          }

          .subcard {
            background: var(--subcard);
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 18px;
            margin-top: 14px;
          }

          .row {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 12px;
          }

          input[type=file],
          input[type=text],
          select,
          textarea {
            width: 100%;
            background: var(--input-bg);
            color: var(--text);
            border: 1px solid var(--border);
            padding: 12px;
            border-radius: 10px;
            box-sizing: border-box;
          }

          textarea {
            min-height: 240px;
            resize: vertical;
            font-family: inherit;
          }

          label {
            font-size: 12px;
            color: var(--muted-soft);
            display: block;
            margin-bottom: 6px;
          }

          fieldset {
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 12px;
          }

          legend {
            color: var(--muted-soft);
            font-size: 12px;
            padding: 0 6px;
          }

          button {
            margin-top: 14px;
            width: 100%;
            padding: 12px 16px;
            border: 0;
            border-radius: 10px;
            background: var(--button);
            color: white;
            font-weight: 600;
            cursor: pointer;
          }

          button:hover {
            background: var(--button-hover);
          }

          button:disabled {
            opacity: 0.7;
            cursor: not-allowed;
          }

          small {
            display: block;
            margin-top: 10px;
            color: var(--muted-soft);
          }

          a {
            color: #6f97ff;
            text-decoration: none;
          }

          .hint {
            margin-top: 10px;
            font-size: 12px;
            color: var(--muted-soft);
          }

          code {
            background: var(--code-bg);
            padding: 2px 6px;
            border-radius: 6px;
          }

          .status {
            margin-top: 12px;
            font-size: 14px;
          }

          .success { color: var(--success); }
          .error { color: var(--error); }

          .mono {
            font-family: ui-monospace, SFMono-Regular, Menlo, Consolas, monospace;
          }

          .stack {
            display: flex;
            gap: 12px;
            align-items: center;
          }

          .hidden {
            display: none;
          }

          .mode-row {
            margin-bottom: 18px;
          }

          .result-box {
            margin-top: 16px;
            padding-top: 6px;
          }

          ul, ol {
            margin: 10px 0 0 18px;
            padding: 0;
            color: var(--muted);
            line-height: 1.6;
          }

          li { margin-bottom: 8px; }

          .progress-wrap {
            margin-top: 14px;
            display: none;
          }

          .progress-wrap.show {
            display: block;
          }

          .progress-label {
            font-size: 13px;
            color: var(--muted);
            margin-bottom: 8px;
          }

          .progress-bar-shell {
            width: 100%;
            height: 12px;
            border-radius: 999px;
            overflow: hidden;
            background: var(--code-bg);
            border: 1px solid var(--border);
          }

          .progress-bar-fill {
            width: 0%;
            height: 100%;
            background: var(--button);
            transition: width 0.25s ease;
          }

          .progress-meta {
            margin-top: 8px;
            display: flex;
            justify-content: space-between;
            gap: 12px;
            font-size: 12px;
            color: var(--muted-soft);
          }

          .preview-group {
            margin-top: 12px;
            padding: 12px;
            border: 1px solid var(--border);
            border-radius: 10px;
            background: var(--subcard);
          }

          .preview-group-title {
            font-weight: 700;
            margin-bottom: 8px;
          }

          @media (max-width: 700px) {
            .row { grid-template-columns: 1fr; }
          }
        </style>
      </head>
      <body>
        <div class="wrap">
          <div class="topbar">
            <button class="theme-toggle" id="theme-toggle" type="button">Toggle theme</button>
          </div>

          <h1>LucidScript</h1>
          <div class="version">Version __APP_VERSION__ • Whisper model: __MODEL_NAME__</div>
          <p>Choose a mode, then process audio, pasted text, or image text into a formatted .docx.</p>
          <div class="card">
            <div class="mode-row">
              <label for="mode">Mode</label>
              <select id="mode">
                <option value="audio">Audio Transcription</option>
                <option value="text">Text Input</option>
                <option value="image">Image Upload</option>
              </select>
            </div>

            <div id="mode-audio">
              <h2>Audio Transcription</h2>
              <p>Upload audio or paste a YouTube URL → optional language/translate → choose output style → download .docx.</p>
              <div class="hint" style="margin-bottom:12px;">
                Pipeline: upload → temporary file → Whisper transcription → optional translation / speaker flow → DOCX export
              </div>

              <form id="ls-form">
  <label for="audio_source">Audio source</label>
  <select id="audio_source" name="audio_source">
    <option value="upload">Upload File</option>
    <option value="youtube">YouTube URL</option>
  </select>

  <div id="audio-upload-group" style="margin-top:12px;">
    <label>Audio file</label>
    <input id="file" type="file" name="file"
           accept=".wav,.mp3,.m4a,.aac,.flac,.ogg,.webm,.mp4,audio/*,video/*" />
  </div>

  <div id="youtube-url-group" class="hidden" style="margin-top:12px;">
    <label for="youtube_url">YouTube URL</label>
    <input
      id="youtube_url"
      type="text"
      name="youtube_url"
      placeholder="Paste a YouTube link here..."
    />
  </div>

  <div class="row" style="margin-top:12px">
    <div>
      <label>Language (choose or Auto)</label>
      <select id="language" name="language">
        <option value="">Auto-detect</option>
        <option value="en">English — en</option>
        <option value="es">Spanish — es</option>
        <option value="pt">Portuguese — pt</option>
        <option value="zh">Mandarin Chinese — zh</option>
        <option value="fr">French — fr</option>
      </select>
    </div>

    <div class="stack" style="margin-top:24px">
      <input type="checkbox" id="translate" name="translate" value="true" />
      <label for="translate" style="margin:0;">Translate to English</label>
    </div>
  </div>

  <div class="row" style="margin-top:12px">
    <fieldset>
      <legend>Output style</legend>
      <div class="stack">
        <input type="radio" id="style-standard" name="style" value="standard" checked />
        <label for="style-standard" style="margin:0;">Standard (paragraph doc)</label>
      </div>
      <div class="stack" style="margin-top:6px">
        <input type="radio" id="style-deposition" name="style" value="deposition" />
        <label for="style-deposition" style="margin:0;">Deposition (Q/A with speaker labels)</label>
      </div>
    </fieldset>

    <fieldset>
      <legend>Speaker detection</legend>
      <div class="stack">
        <input type="checkbox" id="diarize" name="diarize" value="true" />
        <label for="diarize" style="margin:0;">Detect speakers (beta)</label>
      </div>
      <small>Requires ffmpeg; optional HuggingFace token improves labeling.</small>
    </fieldset>
  </div>

  <button id="audio-submit-btn" type="submit">Transcribe & Export</button>
</form>

              <div id="audio-progress-wrap" class="progress-wrap">
  <div id="audio-progress-label" class="progress-label">Preparing upload…</div>
  <div class="progress-bar-shell">
    <div id="audio-progress-fill" class="progress-bar-fill"></div>
  </div>
  <div class="progress-meta">
    <span id="audio-progress-stage">Waiting to start</span>
    <span id="audio-progress-percent">0%</span>
  </div>
  <div id="audio-progress-explainer" class="hint" style="margin-top:10px; text-align:left;">
    Large audio and YouTube jobs can take several minutes. If the bar pauses, LucidScript is usually still transcribing in the background. Please do not refresh or close this page while processing is in progress.
  </div>
</div>

              <div class="hint">
                Supported: <code>WAV</code>, <code>MP3</code>, <code>M4A</code>, <code>AAC</code>, <code>FLAC</code>, <code>OGG</code>, <code>WEBM</code>, <code>MP4</code>
              </div>
            </div>

            <div id="mode-text" class="hidden">
              <h2>Text Input</h2>
              <p>Enter text and export it as your chosen format.</p>

              <form id="security-form">
                <label for="report_text">Report text</label>
                <textarea id="report_text" name="report_text" placeholder="Type or paste your report here..." required></textarea>
                <button type="submit">Generate DOCX</button>
              </form>
            </div>

            <div id="mode-image" class="hidden">
              <h2>Image Upload</h2>
              <p>Upload one or more images, extract multilingual text, optionally translate it to English, and export a single combined .docx.</p>

              <form id="security-image-form">
  <label for="image_files">Image files</label>
  <input
    id="image_files"
    type="file"
    name="image_files"
    accept=".png,.jpg,.jpeg,.webp,.bmp,image/*"
    multiple
  />

  <div style="margin-top:10px; display:flex; gap:10px;">
    <button type="button" id="add-image-btn">Add selected image(s)</button>
    <button type="button" id="clear-image-btn">Clear list</button>
  </div>

  <div id="queued-image-list" class="preview-group" style="margin-top:12px;">
    <div class="preview-group-title">Queued images</div>
    <div id="queued-image-items" class="mono">No images added yet.</div>
  </div>

  <div class="stack" style="margin-top:12px">
    <input type="checkbox" id="translate_image_text" name="translate_to_english" value="true" />
    <label for="translate_image_text" style="margin:0;">Translate extracted text to English</label>
  </div>
  
                <button id="image-submit-btn" type="submit">Extract & Export</button>
              </form>

              <div id="image-progress-wrap" class="progress-wrap">
                <div id="image-progress-label" class="progress-label">Preparing upload…</div>
                <div class="progress-bar-shell">
                  <div id="image-progress-fill" class="progress-bar-fill"></div>
                </div>
                <div class="progress-meta">
                  <span id="image-progress-stage">Waiting to start</span>
                  <span id="image-progress-percent">0%</span>
                </div>
              </div>

              <div class="hint">
                Supported: <code>PNG</code>, <code>JPG</code>, <code>JPEG</code>, <code>WEBP</code>, <code>BMP</code> • Multiple files allowed
              </div>

              <div class="subcard">
                <h3>OCR / Image processing pipeline</h3>
                <ol>
                  <li><strong>Upload:</strong> the user uploads an image file through the interface.</li>
                  <li><strong>Validation:</strong> the system checks that the file is a supported image format.</li>
                  <li><strong>Temporary staging:</strong> the image is saved as a temporary local file for processing.</li>
                  <li><strong>OCR extraction:</strong> the system runs the image through the OCR reader to detect readable text.</li>
                  <li><strong>Fallback language support:</strong> if needed, alternate OCR readers are used for Chinese or Japanese text.</li>
                  <li><strong>Optional translation:</strong> if selected, extracted text is translated into English.</li>
                  <li><strong>Formatting:</strong> the extracted text is cleaned and organized into a readable report structure.</li>
                  <li><strong>DOCX generation:</strong> a Word document is created from the extracted text.</li>
                  <li><strong>Download:</strong> the final document is returned to the user through a download link.</li>
                </ol>
              </div>

              <div class="subcard">
                <h3>OCR considerations</h3>
                <ul>
                  <li>Accuracy depends on image quality, lighting, resolution, and text clarity.</li>
                  <li>Printed text is more reliable than handwriting or highly stylized fonts.</li>
                  <li>Different languages and layouts may require different OCR readers or cleanup steps.</li>
                  <li>Future improvements may include better preprocessing, layout detection, and stronger multilingual support.</li>
                </ul>
              </div>
            </div>

            <div id="shared-status" class="status result-box"></div>
            <div id="shared-result" class="result-box"></div>

            <small>Prefer the API? See <a href="/docs">/docs</a>.</small>
          </div>
        </div>

        <script>
          const root = document.documentElement;
          const modeSelect = document.getElementById('mode');
          const audioPanel = document.getElementById('mode-audio');
          const textPanel = document.getElementById('mode-text');
          const imagePanel = document.getElementById('mode-image');

          const sharedStatusEl = document.getElementById('shared-status');
          const sharedResultEl = document.getElementById('shared-result');

          const audioProgressWrap = document.getElementById('audio-progress-wrap');
          const audioProgressFill = document.getElementById('audio-progress-fill');
          const audioProgressLabel = document.getElementById('audio-progress-label');
          const audioProgressStage = document.getElementById('audio-progress-stage');
          const audioProgressPercent = document.getElementById('audio-progress-percent');
          const audioSubmitBtn = document.getElementById('audio-submit-btn');

          const imageProgressWrap = document.getElementById('image-progress-wrap');
          const imageProgressFill = document.getElementById('image-progress-fill');
          const imageProgressLabel = document.getElementById('image-progress-label');
          const imageProgressStage = document.getElementById('image-progress-stage');
          const imageProgressPercent = document.getElementById('image-progress-percent');
          const imageSubmitBtn = document.getElementById('image-submit-btn');

          const themeToggle = document.getElementById('theme-toggle');

          function applyTheme(theme) {
            root.setAttribute('data-theme', theme);
            localStorage.setItem('lucidscript-theme', theme);
            themeToggle.textContent = theme === 'dark' ? 'Switch to light mode' : 'Switch to dark mode';
          }

          const savedTheme = localStorage.getItem('lucidscript-theme') || 'dark';
          applyTheme(savedTheme);

          themeToggle.addEventListener('click', () => {
            const current = root.getAttribute('data-theme') || 'dark';
            applyTheme(current === 'dark' ? 'light' : 'dark');
          });

          function setMode(mode) {
  audioPanel.classList.add('hidden');
  textPanel.classList.add('hidden');
  imagePanel.classList.add('hidden');

  if (mode === 'audio') {
    audioPanel.classList.remove('hidden');
  } else if (mode === 'text') {
    textPanel.classList.remove('hidden');
  } else if (mode === 'image') {
    imagePanel.classList.remove('hidden');
  }

  sharedStatusEl.className = 'status result-box';
  sharedStatusEl.textContent = '';
  sharedResultEl.innerHTML = '';
}

          modeSelect.addEventListener('change', (e) => {
            setMode(e.target.value);
          });

          function escapeHtml(value) {
            return String(value)
              .replace(/&/g, '&amp;')
              .replace(/</g, '&lt;')
              .replace(/>/g, '&gt;')
              .replace(/"/g, '&quot;')
              .replace(/'/g, '&#039;');
          }

          function showError(message) {
            sharedStatusEl.className = 'status error result-box';
            sharedStatusEl.textContent = message;
          }

          function showSuccess(message) {
            sharedStatusEl.className = 'status success result-box';
            sharedStatusEl.textContent = message;
          }

          function resetProgress(wrap, fill, label, stage, percentEl) {
            wrap.classList.remove('show');
            fill.style.width = '0%';
            label.textContent = 'Preparing upload…';
            stage.textContent = 'Waiting to start';
            percentEl.textContent = '0%';
          }

          function showProgress(wrap) {
  wrap.classList.add('show');
}

function updateAudioSourceUI() {
  const source = audioSourceSelect.value;

  if (source === 'youtube') {
    audioUploadGroup.classList.add('hidden');
    youtubeUrlGroup.classList.remove('hidden');
    audioFileInput.required = false;
    youtubeUrlInput.required = true;
  } else {
    audioUploadGroup.classList.remove('hidden');
    youtubeUrlGroup.classList.add('hidden');
    audioFileInput.required = true;
    youtubeUrlInput.required = false;
  }
}

          function setProgress(fill, label, stage, percentEl, percent, labelText, stageText) {
            const safePercent = Math.max(0, Math.min(100, percent));
            fill.style.width = safePercent + '%';
            label.textContent = labelText;
            stage.textContent = stageText;
            percentEl.textContent = Math.round(safePercent) + '%';
          }

          function makeProcessingSimulator(fill, label, stage, percentEl, steps) {
  let index = 0;
  let interval = null;
  let crawlInterval = null;
  let currentPercent = 0;

  function start() {
    stop();
    interval = setInterval(() => {
      if (index >= steps.length) {
        clearInterval(interval);
        interval = null;
        beginSlowCrawl();
        return;
      }

      const step = steps[index];
      currentPercent = step.percent;
      setProgress(fill, label, stage, percentEl, step.percent, step.label, step.stage);
      index += 1;
    }, 900);
  }

  function beginSlowCrawl() {
    if (crawlInterval) return;

    crawlInterval = setInterval(() => {
      if (currentPercent < 92) {
        currentPercent += 1;
        setProgress(
          fill,
          label,
          stage,
          percentEl,
          currentPercent,
          'Still processing…',
          'Transcribing audio — longer jobs can take several minutes'
        );
      }
    }, 2500);
  }

  function stop(finalPercent = null, finalLabel = null, finalStage = null) {
    if (interval) {
      clearInterval(interval);
      interval = null;
    }
    if (crawlInterval) {
      clearInterval(crawlInterval);
      crawlInterval = null;
    }
    if (finalPercent !== null) {
      currentPercent = finalPercent;
      setProgress(fill, label, stage, percentEl, finalPercent, finalLabel || '', finalStage || '');
    }
  }

  return { start, stop };
}

          function sendWithProgress({
            endpoint,
            formData,
            wrap,
            fill,
            label,
            stage,
            percentEl,
            submitButton,
            uploadLabel,
            processingSteps,
            onSuccess,
            onErrorMessage
          }) {
            return new Promise((resolve, reject) => {
              const xhr = new XMLHttpRequest();
              const simulator = makeProcessingSimulator(fill, label, stage, percentEl, processingSteps);

              showProgress(wrap);
              submitButton.disabled = true;
              setProgress(fill, label, stage, percentEl, 0, 'Preparing upload…', 'Initializing');

              xhr.open('POST', endpoint, true);

              xhr.upload.onprogress = function (event) {
  if (event.lengthComputable) {
    const uploadPercent = Math.min(35, (event.loaded / event.total) * 35);
    setProgress(
      fill,
      label,
      stage,
      percentEl,
      uploadPercent,
      uploadLabel,
      uploadLabel === 'Submitting YouTube link…'
        ? 'Submitting link'
        : 'Uploading file to server'
    );
  } else {
    setProgress(
      fill,
      label,
      stage,
      percentEl,
      10,
      uploadLabel,
      uploadLabel === 'Submitting YouTube link…'
        ? 'Submitting link'
        : 'Uploading file to server'
    );
  }
};

              xhr.onreadystatechange = function () {
                if (xhr.readyState === XMLHttpRequest.HEADERS_RECEIVED && uploadLabel !== 'Submitting YouTube link…') {
  simulator.start();
}

                if (xhr.readyState === XMLHttpRequest.DONE) {
                  simulator.stop();

                  submitButton.disabled = false;

                  let data = {};
                  try {
                    data = JSON.parse(xhr.responseText || '{}');
                  } catch (e) {
                    data = {};
                  }

                  if (xhr.status >= 200 && xhr.status < 300) {
                    setProgress(fill, label, stage, percentEl, 100, 'Complete', 'Finished');
                    onSuccess(data);
                    resolve(data);
                  } else {
                    setProgress(fill, label, stage, percentEl, 100, 'Failed', 'Error');
                    reject(data.detail || onErrorMessage);
                  }
                }
              };

              xhr.onerror = function () {
                simulator.stop(100, 'Failed', 'Network error');
                submitButton.disabled = false;
                reject(onErrorMessage);
              };

              xhr.send(formData);

              // Start progress immediately for YouTube jobs
              if (uploadLabel === 'Submitting YouTube link…') {
              simulator.start();
              }
            });
          }

          function renderImagePreviewGroups(files) {
            if (!Array.isArray(files) || files.length === 0) {
              return '';
            }

            return files.map((file) => {
              const previewText = file.final_text || file.extracted_text || '[No readable text detected]';
              return `
                <div class="preview-group">
                  <div class="preview-group-title">${escapeHtml(file.filename || 'Unnamed image')}</div>
                  <div class="mono" style="white-space:pre-wrap;">${escapeHtml(previewText)}</div>
                </div>
              `;
            }).join('');
          }

          const audioForm = document.getElementById('ls-form');
          const audioSourceSelect = document.getElementById('audio_source');
          const audioUploadGroup = document.getElementById('audio-upload-group');
          const youtubeUrlGroup = document.getElementById('youtube-url-group');
          const audioFileInput = document.getElementById('file');
          const youtubeUrlInput = document.getElementById('youtube_url');
          audioForm.addEventListener('submit', async (e) => {
  e.preventDefault();
  sharedStatusEl.className = 'status result-box';
  sharedResultEl.innerHTML = '';
  sharedStatusEl.textContent = 'Starting audio job…';

  const fd = new FormData(audioForm);
  fd.set('translate', document.getElementById('translate').checked ? 'true' : 'false');
  fd.set('diarize', document.getElementById('diarize').checked ? 'true' : 'false');

  const style = (document.querySelector('input[name="style"]:checked') || {}).value || 'standard';
  const source = audioSourceSelect.value;

  let endpoint = style === 'deposition'
    ? '/export_docx_from_audio_v3'
    : '/export_docx_from_audio_v2';

  if (source === 'youtube') {
    endpoint = style === 'deposition'
      ? '/export_docx_from_youtube_v3'
      : '/export_docx_from_youtube_v2';
  }

  const uploadLabel = source === 'youtube'
    ? 'Submitting YouTube link…'
    : 'Uploading audio…';

  const processingSteps = source === 'youtube'
    ? [
        { percent: 40, label: 'Link received', stage: 'Downloading video audio' },
        { percent: 52, label: 'Audio ready', stage: 'Preparing transcription' },
        { percent: 64, label: 'Transcribing audio…', stage: 'Running Whisper on downloaded audio' },
        { percent: 78, label: 'Still transcribing…', stage: 'Longer videos may take several minutes' },
        { percent: 90, label: 'Formatting document…', stage: 'Generating DOCX' },
        { percent: 98, label: 'Wrapping things up…', stage: 'Preparing download link' }
      ]
    : [
        { percent: 45, label: 'Upload complete', stage: 'Temporary file created' },
        { percent: 60, label: 'Processing audio…', stage: 'Running Whisper transcription' },
        { percent: 75, label: 'Building output…', stage: 'Cleaning / structuring transcript' },
        { percent: 90, label: 'Formatting document…', stage: 'Generating DOCX' },
        { percent: 97, label: 'Finalizing…', stage: 'Preparing download link' }
      ];

  try {
    await sendWithProgress({
      endpoint,
      formData: fd,
      wrap: audioProgressWrap,
      fill: audioProgressFill,
      label: audioProgressLabel,
      stage: audioProgressStage,
      percentEl: audioProgressPercent,
      submitButton: audioSubmitBtn,
      uploadLabel,
      processingSteps,
      onSuccess: (data) => {
        showSuccess('Done – .docx is ready below.');

        const lang = data.language || 'unknown';
        const dur = (data.duration_sec !== null && data.duration_sec !== undefined)
          ? data.duration_sec
          : '—';
        const fname = data.docx_filename;

        sharedResultEl.innerHTML = `
          <div class="mono">
            Language: ${escapeHtml(lang)} |
            Duration: ${escapeHtml(dur)}s |
            Version: ${escapeHtml(data.version || '__APP_VERSION__')}
          </div>
          <div style="margin-top:8px">
            <a href="/download/${encodeURIComponent(fname)}">
              ⬇️ Download ${escapeHtml(fname)}
            </a>
          </div>
        `;
      },
      onErrorMessage: 'Transcription failed. Try a smaller file or shorter clip.'
    });
  } catch (err) {
    showError(String(err || 'Transcription failed.'));
  }
});


          const textForm = document.getElementById('security-form');
          textForm.addEventListener('submit', async (e) => {
            e.preventDefault();
            sharedStatusEl.className = 'status result-box';
            sharedResultEl.innerHTML = '';
            sharedStatusEl.textContent = 'Formatting report…';

            const fd = new FormData(textForm);

            try {
              const resp = await fetch('/export_security_report', { method: 'POST', body: fd });
              const data = await resp.json();

              if (!resp.ok) {
                showError(data.detail || 'Report formatting failed.');
                return;
              }

              showSuccess('Done – .docx is ready below.');

              const fname = data.docx_filename;

              sharedResultEl.innerHTML = `
                <div style="margin-top:8px">
                  <a href="/download/${encodeURIComponent(fname)}">⬇️ Download ${escapeHtml(fname)}</a>
                </div>
              `;
            } catch (err) {
              showError('Unexpected error: ' + (err?.message || err));
            }
          });

          const imageForm = document.getElementById('security-image-form');
const imageFilesInput = document.getElementById('image_files');
const addImageBtn = document.getElementById('add-image-btn');
const clearImageBtn = document.getElementById('clear-image-btn');
const queuedImageItems = document.getElementById('queued-image-items');

let queuedImages = [];

function renderQueuedImages() {
  if (!queuedImages.length) {
    queuedImageItems.textContent = 'No images added yet.';
    return;
  }

  queuedImageItems.innerHTML = queuedImages
    .map((file, index) => {
      return `
        <div style="display:flex; justify-content:space-between; margin-bottom:6px;">
          <span>${escapeHtml(file.name)}</span>
          <button type="button" onclick="removeQueuedImage(${index})">Remove</button>
        </div>
      `;
    })
    .join('');
}

window.removeQueuedImage = function(index) {
  queuedImages.splice(index, 1);
  renderQueuedImages();
};

addImageBtn.addEventListener('click', () => {
  const selectedFiles = Array.from(imageFilesInput.files || []);

  if (!selectedFiles.length) {
    showError('Choose at least one image to add.');
    return;
  }

  selectedFiles.forEach((file) => {
    const exists = queuedImages.some(
      (f) => f.name === file.name && f.size === file.size
    );
    if (!exists) {
      queuedImages.push(file);
    }
  });

  imageFilesInput.value = '';
  renderQueuedImages();
});

clearImageBtn.addEventListener('click', () => {
  queuedImages = [];
  imageFilesInput.value = '';
  renderQueuedImages();
});

imageForm.addEventListener('submit', async (e) => {
  e.preventDefault();
  sharedStatusEl.className = 'status result-box';
  sharedResultEl.innerHTML = '';
  sharedStatusEl.textContent = 'Starting image job…';

  if (!queuedImages.length) {
    showError('Please add at least one image.');
    return;
  }

  const fd = new FormData();

  queuedImages.forEach((file) => {
    fd.append('image_files', file);
  });

  fd.set(
    'translate_to_english',
    document.getElementById('translate_image_text').checked ? 'true' : 'false'
  );

  try {
    await sendWithProgress({
      endpoint: '/export_multi_image_ocr',
      formData: fd,
      wrap: imageProgressWrap,
      fill: imageProgressFill,
      label: imageProgressLabel,
      stage: imageProgressStage,
      percentEl: imageProgressPercent,
      submitButton: imageSubmitBtn,
      uploadLabel: 'Uploading image files…',
      processingSteps: [
        { percent: 45, label: 'Upload complete', stage: 'Temporary images saved' },
        { percent: 62, label: 'Running OCR…', stage: 'Extracting text from uploaded images' },
        { percent: 78, label: 'Processing text…', stage: 'Applying translation / cleanup' },
        { percent: 92, label: 'Formatting document…', stage: 'Generating combined DOCX' },
        { percent: 97, label: 'Finalizing…', stage: 'Preparing download link' }
      ],
      onSuccess: (data) => {
        showSuccess('Done – OCR document is ready below.');

        const fname = data.docx_filename;
        const imageCount = data.image_count || 0;

        sharedResultEl.innerHTML = `
          <div class="mono">Images processed: ${escapeHtml(imageCount)} | Version: ${escapeHtml(data.version || '__APP_VERSION__')}</div>
          <div style="margin-top:8px">
            <a href="/download/${encodeURIComponent(fname)}">⬇️ Download ${escapeHtml(fname)}</a>
          </div>
          <div class="hint" style="margin-top:12px;">Preview by filename:</div>
          ${renderImagePreviewGroups(data.files)}
        `;

        queuedImages = [];
        renderQueuedImages();
      },
      onErrorMessage: 'Multi-image OCR failed.'
    });
  } catch (err) {
    showError(String(err || 'Image OCR failed.'));
  }
});

renderQueuedImages();

          audioSourceSelect.addEventListener('change', updateAudioSourceUI);
updateAudioSourceUI();
setMode('audio');
        </script>
      </body>
    </html>
    """
    return page.replace("__APP_VERSION__", html.escape(APP_VERSION)).replace(
        "__MODEL_NAME__", html.escape(model_name)
    )


@app.get("/download/{filename}")
def download_file(filename: str):
    if "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename.")
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")
    return FileResponse(
        file_path.as_posix(),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/export_docx_from_audio_v2")
async def export_docx_from_audio_v2(
    file: UploadFile = File(...),
    language: str | None = Form(None),
    translate: str | None = Form(None),
):
    suffix = os.path.splitext(file.filename or "")[-1]
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        kwargs = {}
        if language:
            kwargs["language"] = language
        if (translate or "").lower() == "true":
            kwargs["task"] = "translate"

        result = await asyncio.to_thread(get_model().transcribe, tmp_path, **kwargs)
        text = (result.get("text") or "").strip()
        if not text:
            raise HTTPException(
                status_code=400,
                detail="No speech was detected in this file. Try a different recording.",
            )

        lang = result.get("language", "unknown")
        duration = (
            round(float(result.get("duration", 0)), 2) if "duration" in result else None
        )
        translated_flag = (translate or "").lower() == "true"

        out = await asyncio.to_thread(
            build_transcript_doc,
            "LucidScript Transcript",
            text,
            lang,
            translated_flag,
        )

        save_document_record(
            mode="audio",
            original_filename=file.filename,
            output_filename=out.name,
            status="completed",
            language=lang,
            translated=translated_flag,
            notes="Generated from async audio route v2.",
        )

        return JSONResponse(
            {
                "message": "Transcription finished. Word doc ready to download.",
                "docx_path": str(out),
                "docx_filename": out.name,
                "language": lang,
                "duration_sec": duration,
                "translated": translated_flag,
                "version": APP_VERSION,
            }
        )
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Something went wrong while processing the audio. Please try again with a different file.",
        )
    finally:
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass


HUGGINGFACE_TOKEN = os.getenv("HUGGINGFACE_TOKEN")
try:
    from pyannote.audio import Pipeline as PyannotePipeline  # type: ignore

    _PYANNOTE_OK = True
except Exception:
    _PYANNOTE_OK = False


def _time_fmt(t: float):
    t = max(0.0, float(t))
    m = int(t // 60)
    s = int(round(t - 60 * m))
    return f"{m:02d}:{s:02d}"


def _convert_to_wav(src: str):
    out = (OUTPUT_DIR / f"tmp_{uuid.uuid4().hex[:8]}.wav").as_posix()
    cmd = f"ffmpeg -y -i {shlex.quote(src)} -ac 1 -ar 16000 {shlex.quote(out)}"
    try:
        subprocess.run(
            cmd,
            shell=True,
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        return out
    except Exception:
        return src


def download_youtube_audio(url: str) -> tuple[str, dict]:
    """
    Downloads audio from a YouTube URL and returns:
    (local_audio_path, metadata)
    """
    try:
        output_path = (OUTPUT_DIR / f"yt_{uuid.uuid4().hex[:8]}.%(ext)s").as_posix()

        cmd = [
            "yt-dlp",
            "-f",
            "bestaudio",
            "--extract-audio",
            "--audio-format",
            "wav",
            "--audio-quality",
            "0",
            "--no-playlist",
            "--print-json",
            "-o",
            output_path,
            url,
        ]

        result = subprocess.run(cmd, capture_output=True, text=True)

        if result.returncode != 0:
            raise Exception(f"yt-dlp error:\n{result.stderr}")

        metadata = {}
        try:
            metadata = json.loads(result.stdout.splitlines()[-1])
        except Exception:
            pass

        base = output_path.replace(".%(ext)s", "")
        for ext in [".wav", ".mp3", ".m4a"]:
            candidate = base + ext
            if os.path.exists(candidate):
                return candidate, metadata

        raise Exception("Audio file not found after download.")

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"YouTube download failed: {str(e)}"
        )


async def process_audio_file_to_docx(
    audio_path: str,
    original_name: str | None = None,
    language: str | None = None,
    translate: bool = False,
    notes: str | None = None,
):
    kwargs = {}
    if language:
        kwargs["language"] = language
    if translate:
        kwargs["task"] = "translate"

    result = await asyncio.to_thread(get_model().transcribe, audio_path, **kwargs)
    text = (result.get("text") or "").strip()

    if not text:
        raise HTTPException(
            status_code=400,
            detail="No speech was detected in this file. Try a different recording.",
        )

    detected_language = result.get("language", "unknown")
    duration = (
        round(float(result.get("duration", 0)), 2) if "duration" in result else None
    )

    out = await asyncio.to_thread(
        build_transcript_doc,
        "LucidScript Transcript",
        text,
        detected_language,
        translate,
    )

    save_document_record(
        mode="audio",
        original_filename=original_name,
        output_filename=out.name,
        status="completed",
        language=detected_language,
        translated=translate,
        notes=notes or "Generated from shared audio processing helper.",
    )

    return {
        "docx_path": str(out),
        "docx_filename": out.name,
        "language": detected_language,
        "duration_sec": duration,
        "translated": translate,
        "version": APP_VERSION,
    }


def _diarize_segments_dep(wav_path: str) -> List[Tuple[float, float, str]]:
    if not (_PYANNOTE_OK and HUGGINGFACE_TOKEN):
        return []
    try:
        pipe = PyannotePipeline.from_pretrained(
            "pyannote/speaker-diarization-3.1",
            use_auth_token=HUGGINGFACE_TOKEN,
        )
        diar = pipe(wav_path)
        segs = []
        for turn, _, spk in diar.itertracks(yield_label=True):
            segs.append((float(turn.start), float(turn.end), str(spk)))
        segs.sort(key=lambda x: x[0])
        return segs
    except Exception:
        return []


def _assign_speakers(segments: List[dict], dia: List[Tuple[float, float, str]]):
    if not dia:
        return [
            {
                "speaker": "Speaker 1",
                "start": s.get("start", 0.0),
                "end": s.get("end", 0.0),
                "text": (s.get("text") or "").strip(),
            }
            for s in segments
        ]
    labeled = []
    for seg in segments:
        s = seg.get("start", 0.0)
        e = seg.get("end", 0.0)
        txt = (seg.get("text") or "").strip()
        best, overlap = "Speaker 1", 0.0
        for ds, de, spk in dia:
            ov = max(0.0, min(e, de) - max(s, ds))
            if ov > overlap:
                best, overlap = spk, ov
        labeled.append({"speaker": best, "start": s, "end": e, "text": txt})
    return labeled


def _make_deposition_doc(
    title: str, language: str, translated: bool, labeled: List[dict]
):
    doc = Document()
    doc.add_heading(title, 0)
    meta = f"{datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Language: {language or 'unknown'}"
    if translated:
        meta += "  |  Translated→English"
    doc.add_paragraph(meta)

    line_limit = 25
    current_line = 0
    for seg in labeled:
        header = (
            f"{seg['speaker']}  [{_time_fmt(seg['start'])}–{_time_fmt(seg['end'])}]"
        )
        p = doc.add_paragraph(header)
        p.runs[0].bold = True
        for line in seg["text"].splitlines() or [""]:
            wrapped = "\n".join(textwrap.wrap(line, width=80)) or ""
            for sub in wrapped.split("\n") if wrapped else [""]:
                if current_line >= line_limit:
                    doc.add_page_break()
                    current_line = 0
                doc.add_paragraph(f"    {sub}")
                current_line += 1

    out = OUTPUT_DIR / f"lucidscript_{uuid.uuid4().hex[:8]}.docx"
    doc.save(out.as_posix())
    return out


@app.post("/export_docx_from_audio_v3")
async def export_docx_from_audio_v3(
    file: UploadFile = File(...),
    language: str | None = Form(None),
    translate: str | None = Form(None),
    diarize: str | None = Form(None),
):
    suffix = os.path.splitext(file.filename or "")[-1]
    tmp_path = None
    wav16k = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        kwargs = {}
        if language:
            kwargs["language"] = language
        if (translate or "").lower() == "true":
            kwargs["task"] = "translate"

        result = await asyncio.to_thread(get_model().transcribe, tmp_path, **kwargs)
        text = (result.get("text") or "").strip()
        if not text:
            raise HTTPException(
                status_code=400,
                detail="No speech was detected in this file. Try a different recording.",
            )

        segments = result.get("segments", [])
        do_diar = (diarize or "").lower() == "true"
        if do_diar:
            wav16k = await asyncio.to_thread(_convert_to_wav, tmp_path)
            dia = await asyncio.to_thread(_diarize_segments_dep, wav16k)
            labeled = _assign_speakers(segments, dia)
        else:
            labeled = _assign_speakers(segments, [])

        out = await asyncio.to_thread(
            _make_deposition_doc,
            "LucidScript Deposition Transcript",
            result.get("language", "unknown"),
            ((translate or "").lower() == "true"),
            labeled,
        )

        save_document_record(
            mode="audio",
            original_filename=file.filename,
            output_filename=out.name,
            status="completed",
            language=result.get("language", "unknown"),
            translated=((translate or "").lower() == "true"),
            notes="Generated from deposition audio route v3.",
        )

        return JSONResponse(
            {
                "message": "Deposition transcript complete. Word doc ready to download.",
                "docx_path": str(out),
                "docx_filename": out.name,
                "language": result.get("language", "unknown"),
                "duration_sec": (
                    round(float(result.get("duration", 0)), 2)
                    if "duration" in result
                    else None
                ),
                "translated": ((translate or "").lower() == "true"),
                "version": APP_VERSION,
            }
        )
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="Something went wrong while processing the audio. Please try a different file.",
        )
    finally:
        try:
            if wav16k and wav16k != tmp_path and os.path.exists(wav16k):
                os.remove(wav16k)
        except Exception:
            pass
        try:
            if tmp_path and os.path.exists(tmp_path):
                os.remove(tmp_path)
        except Exception:
            pass


@app.post("/export_docx_from_youtube_v2")
async def export_docx_from_youtube_v2(
    youtube_url: str = Form(...),
    language: str | None = Form(None),
    translate: bool = Form(False),
    diarize: bool = Form(False),
):
    audio_path = None
    try:
        # Download audio
        audio_path, metadata = download_youtube_audio(youtube_url)
        title = metadata.get("title") or "youtube_audio"

        # Process
        result = await process_audio_file_to_docx(
            audio_path=audio_path,
            original_name=title,
            language=language,
            translate=translate,
            notes=f"YouTube source: {youtube_url}",
        )

        return result

    finally:
        if audio_path and os.path.exists(audio_path):
            try:
                os.remove(audio_path)
            except Exception:
                pass


@app.post("/export_docx_from_youtube_v3")
async def export_docx_from_youtube_v3(
    youtube_url: str = Form(...),
    language: str | None = Form(None),
    translate: str | None = Form(None),
    diarize: str | None = Form(None),
):
    audio_path = None
    wav_path = None
    try:
        translate_flag = (translate or "").lower() == "true"
        diarize_flag = (diarize or "").lower() == "true"

        audio_path, metadata = download_youtube_audio(youtube_url)
        title = metadata.get("title") or "youtube_audio"

        wav_path = _convert_to_wav(audio_path)

        kwargs = {}
        if language:
            kwargs["language"] = language
        if translate_flag:
            kwargs["task"] = "translate"

        result = await asyncio.to_thread(get_model().transcribe, wav_path, **kwargs)
        text = (result.get("text") or "").strip()

        if not text:
            raise HTTPException(status_code=400, detail="No speech detected.")

        detected_language = result.get("language", "unknown")

        if diarize_flag:
            dia = await asyncio.to_thread(_diarize_segments_dep, wav_path)
            labeled = _assign_speakers(result.get("segments", []), dia)
        else:
            labeled = _assign_speakers(result.get("segments", []), [])

        out = await asyncio.to_thread(
            _make_deposition_doc,
            "LucidScript Deposition",
            detected_language,
            translate_flag,
            labeled,
        )

        save_document_record(
            mode="audio",
            original_filename=title,
            output_filename=out.name,
            status="completed",
            language=detected_language,
            translated=translate_flag,
            notes=f"YouTube source: {youtube_url}",
        )

        return {
            "docx_filename": out.name,
            "language": detected_language,
            "translated": translate_flag,
            "version": APP_VERSION,
        }

    finally:
        try:
            if wav_path and wav_path != audio_path and os.path.exists(wav_path):
                os.remove(wav_path)
        except Exception:
            pass
        try:
            if audio_path and os.path.exists(audio_path):
                os.remove(audio_path)
        except Exception:
            pass
